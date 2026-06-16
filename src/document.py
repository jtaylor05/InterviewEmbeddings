from docx import Document

class InterviewDocument:
    def __init__(self, file_path : str):
        self.paragraphs = []
        self._convert_to_paragraphs(file_path)
    
    def _convert_to_paragraphs(self, file_path) -> list[str]:
        if file_path.endswith(".txt"):
            with open(file_path, 'r') as f:
                self.paragraphs = [l for l in f.readlines() if l.strip()]
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            self.paragraphs = [s for p in doc.paragraphs if p.text.strip() for s in p.text.split("\n") if s.strip()]