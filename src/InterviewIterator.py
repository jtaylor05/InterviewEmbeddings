from docx import Document

class InterviewIterator:
    def __init__(self, file_path):
        self.document = Document(file_path)
        self.paragraphs = self.document.paragraphs
        self.current_index = 0

    def get_paragraphs(self):
        return [paragraph.text for paragraph in self.paragraphs]
    
    def __iter__(self):
        return self

    def __next__(self):
        if self.current_index < len(self.paragraphs):
            paragraph = self.paragraphs[self.current_index]
            self.current_index += 1
            return paragraph.text
        else:
            raise StopIteration()